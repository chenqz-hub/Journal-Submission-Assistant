from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
import os
import re

def format_docx(file_path: str, rules_input: dict) -> str:
    """
    Applies formatting rules to a DOCX file.
    """
    try:
        if file_path.lower().endswith(".md") or file_path.lower().endswith(".txt"):
            doc = Document()
            with open(file_path, "r", encoding="utf-8") as f:
                for line in f:
                    if line.strip():
                        doc.add_paragraph(line.strip())
        else:
            doc = Document(file_path)
    except Exception as e:
        # Better error handling for file reading
        raise ValueError(f"无法打开或解析 Word 文档，请确认文件未损坏。({str(e)})")
    
    # Extract rules: Handle nested "data" -> "rules" structure or direct rules
    # The frontend might pass the whole response object
    rules = {}
    if "data" in rules_input and "rules" in rules_input["data"]:
        rules = rules_input["data"]["rules"]
    elif "rules" in rules_input:
        rules = rules_input["rules"]
    else:
        rules = rules_input

    # --- 1. Font Settings ---
    # Default to Times New Roman if not found or if null
    font_name = rules.get("font_family")
    if not font_name or "null" in str(font_name).lower():
        font_name = "Times New Roman"
        
    # Simplify font name (e.g. "Times New Roman, 12pt" -> "Times New Roman")
    if "," in font_name:
        font_name = font_name.split(",")[0].strip()

    # Font Size
    font_size_val = rules.get("font_size_pt")
    if font_size_val is None:
        font_size_val = 12
    font_size = 12 # Default
    try:
        font_size = float(font_size_val)
    except:
        font_size = 12

    # --- 2. Line Spacing Logic ---
    line_spacing_rule = rules.get("line_spacing")
    spacing_val = 1.0 # Default single
    
    if line_spacing_rule:
        ls_str = str(line_spacing_rule).lower()
        if "double" in ls_str or "2.0" in ls_str or "2" == ls_str:
            spacing_val = 2.0
        elif "1.5" in ls_str:
            spacing_val = 1.5
        elif "single" in ls_str or "1.0" in ls_str or "1" == ls_str:
            spacing_val = 1.0

    # --- 3. Apply to Styles (Global) ---
    # Updating 'Normal' style affects most text unless manually overridden
    try:
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
    except:
        pass

    # --- 4. Iterative content formatting ---
    for para in doc.paragraphs:
        # Skip empty paragraphs to avoid clutter? No, keep structure.
        
        # Apply Line Spacing
        # Note: line_spacing_rule = WD_LINE_SPACING.DOUBLE is better than set value 2.0 directly for compatibility
        if spacing_val == 2.0:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.DOUBLE
        elif spacing_val == 1.5:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        else:
            para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        
        # Apply Font to Runs (overwrites manual formatting)
        for run in para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)

    # Output directory handling
    output_dir = os.path.join(os.getcwd(), "static", "output")
    os.makedirs(output_dir, exist_ok=True)
    
    filename = os.path.basename(file_path)
    # Ensure the output is saved as a docx file even if the input was md/txt
    if filename.lower().endswith(('.md', '.txt')):
        filename = os.path.splitext(filename)[0] + ".docx"
        
    return output_filename
