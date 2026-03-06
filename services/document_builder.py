import os
import zipfile
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
import logging

logger = logging.getLogger(__name__)

class DocumentBuilder:
    """
    V2.0 核心模块：学术文档按需打包引擎。
    接收 DocumentSplitter 拆分出的 sections（字典），
    结合期刊的排版规则（如是否盲审，是否需要分离图表等），
    生成多个相应的 DOCX 文件并最终打成一个 ZIP 包供用户下载。
    """
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        
    def _apply_formatting(self, doc: Document, rules: dict):
        """将期刊的字体和行距规则应用到特定的 doc 上"""
        font_name = rules.get("font_family", "Times New Roman")
        if not font_name or "null" in str(font_name).lower():
            font_name = "Times New Roman"
        if "," in font_name:
            font_name = font_name.split(",")[0].strip()

        font_size_val = rules.get("font_size_pt")
        if font_size_val is None:
            font_size_val = 12
        try:
            font_size = float(font_size_val)
        except:
            font_size = 12
            
        line_spacing_val = rules.get("line_spacing", "").lower() if rules.get("line_spacing") else ""

        # Apply to global style
        try:
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = Pt(font_size)
            
            # Apply line spacing
            if "double" in line_spacing_val:
                style.paragraph_format.line_spacing = 2.0
            elif "1.5" in line_spacing_val:
                style.paragraph_format.line_spacing = 1.5
            elif "single" in line_spacing_val:
                style.paragraph_format.line_spacing = 1.0
                
        except:
            pass

        # Also iterate to ensure runs are forced
        for para in doc.paragraphs:
            # Force paragraph line spacing
            if "double" in line_spacing_val:
                para.paragraph_format.line_spacing = 2.0
            elif "1.5" in line_spacing_val:
                para.paragraph_format.line_spacing = 1.5
            elif "single" in line_spacing_val:
                para.paragraph_format.line_spacing = 1.0
                
            for run in para.runs:
                run.font.name = font_name
                run.font.size = Pt(font_size)

    def _set_three_line_table(self, table):
        """将表格转换为学术标准的“三线表”格式"""
        def set_border(tc, top=False, bottom=False, top_sz="12", bottom_sz="12"):
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            
            if top:
                top_border = OxmlElement('w:top')
                top_border.set(qn('w:val'), 'single')
                top_border.set(qn('w:sz'), top_sz) # Border size (1/8 pt, so 12=1.5pt)
                top_border.set(qn('w:space'), '0')
                top_border.set(qn('w:color'), 'auto')
                tcBorders.append(top_border)
            
            if bottom:
                bottom_border = OxmlElement('w:bottom')
                bottom_border.set(qn('w:val'), 'single')
                bottom_border.set(qn('w:sz'), bottom_sz)
                bottom_border.set(qn('w:space'), '0')
                bottom_border.set(qn('w:color'), 'auto')
                tcBorders.append(bottom_border)
                
            tcPr.append(tcBorders)

        # Clear existing borders by setting style to simple list/grid if possible
        table.style = 'Normal Table'
        
        # Apply strict three-line standard
        for r_idx, row in enumerate(table.rows):
            is_first_row = r_idx == 0
            is_last_row = r_idx == len(table.rows) - 1
            
            for c_idx, cell in enumerate(row.cells):
                tc = cell._tc
                if is_first_row:
                    set_border(tc, top=True, bottom=True, top_sz="12", bottom_sz="6") # Top thicker
                elif is_last_row:
                    set_border(tc, top=False, bottom=True, bottom_sz="12") # Bottom thicker
                else:
                    # Middle rows: no borders!
                    pass

    def _add_content(self, doc: Document, item):
        """Helper to properly add a string or a rich text dictionary to a document."""
        if isinstance(item, dict):
            para = doc.add_paragraph()
            style_name = item.get('style')
            if style_name:
                try:
                    para.style = doc.styles[style_name]
                except:
                    pass
            for r_data in item.get('runs', []):
                run = para.add_run(r_data.get('text', ''))
                if r_data.get('bold'): run.bold = True
                if r_data.get('italic'): run.italic = True
                if r_data.get('underline'): run.underline = True
        else:
            doc.add_paragraph(str(item))

    def _create_docx(self, paragraphs: list, filename: str, rules: dict) -> str:
        """从段落列表创建一个格式化好的 docx 文件"""
        doc = Document()
        for p in paragraphs:
            self._add_content(doc, p)

        self._apply_formatting(doc, rules)

        file_path = os.path.join(self.output_dir, filename)
        doc.save(file_path)
        return file_path

    def build_submission_package(self, base_filename: str, sections: dict, rules: dict) -> str:
        """
        核心方法：根据规则组装压缩包
        返回生成的 ZIP 文件的绝对路径。
        """
        # Determine rules with defaults
        is_blind = rules.get("double_blind_review", False)
        # 很多期刊即使不填，也默认建议把图注和表格单拆
        separate_figs = rules.get("figures_separate", True) 
        
        generated_files = []
        name_no_ext = os.path.splitext(base_filename)[0]

        # --- 1. 生成 Title Page ---
        # 无论是否盲审，都生成一个完整的 Title Page 备用，包含作者信息、声明致谢等
        title_page_content = []
        title_page_content.extend(sections.get("title_page", []))
        
        # 将声明和致谢也贴到 title page 中
        if sections.get("acknowledgements"):
            title_page_content.append("\nAcknowledgements")
            title_page_content.extend(sections.get("acknowledgements"))
        if sections.get("declarations"):
            title_page_content.append("\nDeclarations")
            title_page_content.extend(sections.get("declarations"))

        if title_page_content:
            tp_path = self._create_docx(title_page_content, f"Title_Page_{name_no_ext}.docx", rules)
            generated_files.append(tp_path)

        # --- 2. 生成 Main Manuscript ---
        main_content = []
        # 如果不是盲审，正文里可以保留标题甚至部分 title page 内容（出于简单，正文重新从 Abstract 开始）
        # 如果是盲审，必须确保 main content 没有致谢和声明
        if sections.get("abstract"):
            main_content.append("Abstract")
            main_content.extend(sections.get("abstract"))
            
        if sections.get("main_text"):
            # 在组装正文时插入图表占位符（简易实现，实际可以在 splitter 截断的位置直接塞入）
            main_content.extend(sections.get("main_text"))
            
        if not is_blind:
            # 如果不是盲审，可以把致谢放回正文尾部
            if sections.get("acknowledgements"):
                main_content.append("Acknowledgements")
                main_content.extend(sections.get("acknowledgements"))
                
        if sections.get("references"):
            main_content.append("References")
            main_content.extend(sections.get("references"))

        if main_content:
            suffix = "_Blinded" if is_blind else ""
            manuscript_path = self._create_docx(main_content, f"Manuscript{suffix}_{name_no_ext}.docx", rules)
            generated_files.append(manuscript_path)

        # --- 3. 生成 Figure Legends ---
        fig_content = sections.get("figure_legends", [])
        if fig_content:
            fig_doc = Document()
            for fig in fig_content:
                self._add_content(fig_doc, fig)
            self._apply_formatting(fig_doc, rules)
            fig_path = os.path.join(self.output_dir, f"Figure_Legends_{name_no_ext}.docx")
            fig_doc.save(fig_path)
            generated_files.append(fig_path)

        # --- 4. 生成 Tables ---
        # 转移真实的 Table 对象
        tables_obj = sections.get("tables_obj", [])
        if tables_obj:
            tbl_doc = Document()
            # 规矩: 表格一般要比正文小 1-2 磅
            font_size_pt = rules.get("font_size_pt", 12)
            if font_size_pt is None:
                font_size_pt = 12
            try:
                tbl_font_size = float(font_size_pt) - 2
            except (ValueError, TypeError):
                tbl_font_size = 10
            if tbl_font_size < 9: tbl_font_size = 9 # 保底
            
            for i, tbl_group in enumerate(tables_obj):

            
                if i > 0:

            
                    tbl_doc.add_page_break()


            
                if isinstance(tbl_group, dict):

            
                    for t_para in tbl_group.get('title', []):

            
                        self._add_content(tbl_doc, t_para)

            
                    tbl = tbl_group.get('table')

            
                else:

            
                    tbl_doc.add_paragraph(f'Table {i+1}')

            
                    tbl = tbl_group


            
                if not tbl: continue

            
                new_tbl = tbl_doc.add_table(rows=len(tbl.rows), cols=len(tbl.columns))

            
                for r_idx, row in enumerate(tbl.rows):
                    for c_idx, cell in enumerate(row.cells):
                        try:
                            new_cell = new_tbl.cell(r_idx, c_idx)
                            new_cell.text = cell.text
                            new_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                            for p in new_cell.paragraphs:
                                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in p.runs:
                                    run.font.name = rules.get('font_family', 'Times New Roman')
                                    run.font.size = Pt(tbl_font_size)
                        except Exception:
                            pass
                self._set_three_line_table(new_tbl)

            
                if isinstance(tbl_group, dict):

            
                    for f_para in tbl_group.get('footnotes', []):

            
                        self._add_content(tbl_doc, f_para)
                
            tbl_path = os.path.join(self.output_dir, f"Tables_{name_no_ext}.docx")
            tbl_doc.save(tbl_path)
            generated_files.append(tbl_path)

        # --- 5. 打包为 ZIP ---
        zip_filename = f"Submission_Package_{name_no_ext}.zip"
        zip_path = os.path.join(self.output_dir, zip_filename)
        
        with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for file in generated_files:
                # 只保留文件名打包
                zipf.write(file, os.path.basename(file))
                
        return zip_filename

    def build_tables_only(self, base_filename: str, sections: dict, rules: dict) -> str:
        """
        专门生成单独的表格文件。
        """
        import os
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
        name_no_ext = os.path.splitext(base_filename)[0]
        tables_obj = sections.get("tables_obj", [])
        if not tables_obj:
            raise ValueError("No tables found in logic.")

        tbl_doc = Document()
        font_size_pt = rules.get("font_size_pt", 12)
        if font_size_pt is None:
            font_size_pt = 12
        try:
            tbl_font_size = float(font_size_pt) - 2
        except (ValueError, TypeError):
            tbl_font_size = 10
        if tbl_font_size < 9: tbl_font_size = 9 

        for i, tbl_group in enumerate(tables_obj):
            if i > 0:
                tbl_doc.add_page_break()

            if isinstance(tbl_group, dict):
                for t_para in tbl_group.get('title', []):
                    self._add_content(tbl_doc, t_para)
                tbl = tbl_group.get('table')
            else:
                tbl_doc.add_paragraph(f'Table {i+1}')
                tbl = tbl_group

            if not tbl: continue

            new_tbl = tbl_doc.add_table(rows=len(tbl.rows), cols=len(tbl.columns))

            for r_idx, row in enumerate(tbl.rows):
                for c_idx, cell in enumerate(row.cells):
                    try:
                        new_cell = new_tbl.cell(r_idx, c_idx)
                        new_cell.text = cell.text
                        new_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                        for p in new_cell.paragraphs:
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in p.runs:
                                run.font.name = rules.get('font_family', 'Times New Roman')
                                run.font.size = Pt(tbl_font_size)
                    except Exception:
                        pass

            self._set_three_line_table(new_tbl)

            if isinstance(tbl_group, dict):
                for f_para in tbl_group.get('footnotes', []):
                    self._add_content(tbl_doc, f_para)

        tbl_path = os.path.join(self.output_dir, f"TablesOnly_{name_no_ext}.docx")
        tbl_doc.save(tbl_path)
        return f"TablesOnly_{name_no_ext}.docx"
