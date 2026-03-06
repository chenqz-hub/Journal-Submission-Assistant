from fastapi import APIRouter, File, UploadFile, Form, HTTPException
from starlette.concurrency import run_in_threadpool
from pydantic import BaseModel
from typing import Optional
import requests
import json
import logging

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

from services.ai_parser import parse_journal_rules, parse_journal_rules_bundle
from services.web_crawler import fetch_guideline_bundle
from services.pdf_handler import extract_text_from_pdf

router = APIRouter()

class JournalUrlRequest(BaseModel):
    url: str


class GuidelineTextRequest(BaseModel):
    text: str
    source_url: Optional[str] = None
    title: Optional[str] = None

class CoverLetterRequest(BaseModel):
    journal_name: Optional[str] = None
    editor_name: Optional[str] = "Editor"
    guideline_url: Optional[str] = None

@router.post("/parse-rules-pdf")
async def parse_journal_rules_pdf(file: UploadFile = File(...)):
    """
    Upload a PDF file of Author Guidelines -> AI extracts rules.
    """
    if not file.filename.lower().endswith('.pdf'):
        raise HTTPException(status_code=400, detail="File must be a PDF")
    
    try:
        content = await file.read()
        logger.info(f"Received PDF: {file.filename}, Size: {len(content)} bytes")
        
        # Run CPU-bound text extraction in verify threadpool
        text = await run_in_threadpool(extract_text_from_pdf, content)
        
        if not text:
            # Check if text is empty - could be scanned PDF
            logger.warning(f"Extracted empty text from {file.filename}")
            raise HTTPException(status_code=400, detail="Could not extract text from PDF (possibly image-based/scanned). Please use OCR or copy text manually.")
            
        logger.info(f"Extracted text length: {len(text)}")

        # Structure data for parser
        page_data = {
            "text": text,
            "title": file.filename,
            "url": "uploaded_pdf",
        }
        
        # PDF mode should use LLM if needed
        logger.info("Parsing rules with AI...")
        rules = await parse_journal_rules(page_data, use_llm=True)
        return {"status": "success", "message": "PDF rules parsed successfully", "data": rules}

    except HTTPException as he:
        raise he
    except Exception as exc:
        logger.error(f"PDF parsing error: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"PDF parsing failed: {str(exc)}")

@router.post("/parse-rules")
async def parse_rules_endpoint(request: JournalUrlRequest):
    """
    输入期刊「作者指南」链接 → AI 自动抓取、提取格式规则。
    """
    try:
        bundle_data = fetch_guideline_bundle(request.url)
        # For remote fetched data, use sync regex bundle parser mostly
        # because running LLM on 20 candidates sucks for latency.
        # But we could run LLM on the MAIN page result.
        
        parsed_rules = await parse_journal_rules_bundle(bundle_data) 
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=f"输入参数错误: {exc}") from exc
    except PermissionError as exc:
        raise HTTPException(
            status_code=403,
            detail=f"目标网站拒绝访问（已自动尝试多个候选链接）: {exc}",
        ) from exc
    except RuntimeError as exc:
        raise HTTPException(status_code=404, detail=f"未找到可访问的作者指南页面: {exc}") from exc
    except requests.HTTPError as exc:
        status_code = exc.response.status_code if exc.response is not None else 502
        if status_code == 403:
            raise HTTPException(
                status_code=403,
                detail="目标网站拒绝访问（403）。请使用可公开访问的作者指南页面链接，或更换同一期刊的 for-authors/submission 页面。",
            ) from exc
        raise HTTPException(status_code=502, detail=f"规则解析失败: {exc}") from exc
    except Exception as exc:
        raise HTTPException(status_code=502, detail=f"规则解析失败: {exc}") from exc

    return {"status": "success", "message": "解析规则成功", "data": parsed_rules}


@router.post("/parse-rules-text")
async def parse_rules_from_text(request: GuidelineTextRequest):
    """
    粘贴作者指南文本，直接解析格式规则（用于网站反爬时的兜底方案）。
    """
    if not request.text or len(request.text.strip()) < 80:
        raise HTTPException(status_code=400, detail="文本过短，请粘贴完整作者指南内容（至少 80 字符）")

    page_data = {
        "url": request.source_url or "manual-input",
        "title": request.title or "Pasted Guideline Text",
        "fetched_at": None,
        "headings": [],
        "text": request.text.strip(),
    }
    # Text mode should use LLM if needed 
    parsed = await parse_journal_rules(page_data, use_llm=True)
    parsed["crawl"] = {
        "mode": "manual-text",
        "main_url": request.source_url or "manual-input",
        "subpage_count": 0,
        "pages": [],
        "visited_urls": [],
        "main_attempts": [],
        "main_candidates": [],
    }
    return {"status": "success", "message": "文本解析成功", "data": parsed}

@router.post("/format-document")
async def format_document(
    file: UploadFile = File(...), 
    journal_url: Optional[str] = Form(None),
    parsed_rules: Optional[str] = Form(None)
):
    """
    V2.0: 智能文档拆分与打包 (Upload -> Semantic Split -> Build ZIP)
    """
    import json
    import os
    from services.document_splitter import DocumentSplitter
    from services.document_builder import DocumentBuilder

    rules = {}
    if parsed_rules:
        try:
            rules = json.loads(parsed_rules)
            # handle nested structure if frontend sends {data: {rules: {...}}}
            if "data" in rules and "rules" in rules["data"]:
                rules = rules["data"]["rules"]
            elif "rules" in rules:
                rules = rules["rules"]
        except:
            pass
    
    if not rules and journal_url:
        try:
            from services.web_crawler import fetch_guideline_bundle
            from services.ai_parser import parse_journal_rules_bundle
            bundle_data = fetch_guideline_bundle(journal_url)
            rules = parse_journal_rules_bundle(bundle_data)
        except Exception:
            pass

    if not rules:
         raise HTTPException(status_code=400, detail="缺少期刊规则数据，请先解析规则或提供有效URL")

    # Save uploaded file
    content = await file.read()
    temp_path = f"temp/{file.filename}"
    with open(temp_path, "wb") as f:
        f.write(content)
    
    try:
        # V2.0 Workflow: Split -> Build package
        # 1. Semantic Split
        splitter = DocumentSplitter()
        sections = splitter.parse(temp_path)
        
        # 2. Build Docx & Zip
        output_dir = os.path.join(os.getcwd(), "static", "output")
        builder = DocumentBuilder(output_dir)
        zip_filename = builder.build_submission_package(file.filename, sections, rules)
        
        return {
            "status": "success", 
            "message": "文档解构与多文件排版完成，即将下载投稿压缩包", 
            "data": {
                "filename": zip_filename, 
                "download_url": f"/static/output/{zip_filename}"
            }
        }
    except Exception as e:
        import traceback
        logger.error(f"Format pipeline error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"处理失败: {str(e)}")


from typing import List

@router.post("/transfer-document")
async def transfer_document(
    files: List[UploadFile] = File(...), 
    journal_url: Optional[str] = Form(None),
    parsed_rules: Optional[str] = Form(None)
):
    """
    V2.5: 拒稿转投 (Multiple Uploads -> Merge -> Semantic Split/Clean -> Build ZIP)
    """
    import json
    import os
    import shutil
    from services.document_merger import DocumentMerger
    from services.document_builder import DocumentBuilder

    rules = {}
    if parsed_rules:
        try:
            rules = json.loads(parsed_rules)
            if "data" in rules and "rules" in rules["data"]:
                rules = rules["data"]["rules"]
            elif "rules" in rules:
                rules = rules["rules"]
        except:
            pass
    
    if not rules and journal_url:
        try:
            from services.web_crawler import fetch_guideline_bundle
            from services.ai_parser import parse_journal_rules_bundle
            bundle_data = fetch_guideline_bundle(journal_url)
            rules = parse_journal_rules_bundle(bundle_data)
        except Exception:
            pass

    if not rules:
         raise HTTPException(status_code=400, detail="缺少新期刊规则数据，请先解析目标期刊规则")

    temp_paths = []
    try:
        # Save all uploaded files
        for file in files:
            content = await file.read()
            temp_path = f"temp/{file.filename}"
            with open(temp_path, "wb") as f:
                f.write(content)
            temp_paths.append(temp_path)
            
        # Merge Files into single structure
        merger = DocumentMerger()
        sections = merger.merge_files(temp_paths)

        # [NEW] 批量检查图片 DPI
        from services.image_checker import ImageChecker
        dpi_warnings = ImageChecker.check_dpi(temp_paths)

        # [NEW] 参考文献格式转换 (AI 驱动)
        ref_style = rules.get("reference_style")
        if ref_style and str(ref_style).lower() not in ["none", "null", "false", ""] and sections.get("references"):
            logger.info(f"Reformatting references to {ref_style} via LLM...")
            from services.llm_helper import reformat_references_with_llm
            raw_refs_text = "\n".join([r.get('text', '') if isinstance(r, dict) else str(r) for r in sections["references"]])
            if len(raw_refs_text.strip()) > 20: 
                formatted_refs_text = await reformat_references_with_llm(raw_refs_text, str(ref_style))
                # 重新写入转换后的列表
                sections["references"] = [line for line in formatted_refs_text.split('\n') if line.strip()]

        # Build Docx & Zip
        output_dir = os.path.join(os.getcwd(), "static", "output")
        builder = DocumentBuilder(output_dir)
        zip_filename = builder.build_submission_package("Transferred_Submission", sections, rules)

        msg = "多文件合并与新期刊排版完成，即将下载"
        if dpi_warnings:
            msg += "\n(警告：检测到部分图片 DPI 低于 300，解压后请查看详细警告)"

        return {
            "status": "success",
            "message": msg,
            "data": {
                "filename": zip_filename,
                "download_url": f"/static/output/{zip_filename}",
                "dpi_warnings": dpi_warnings
            }
        }
    except Exception as e:
        import traceback
        logger.error(f"Transfer pipeline error: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"多文件转投处理失败: {str(e)}")
    finally:
        # 清理临时文件
        for p in temp_paths:
            if os.path.exists(p):
                try:
                    os.remove(p)
                except Exception as e:
                    logger.warning(f"Failed to remove temp file {p}: {e}")

@router.post("/generate-cover-letter")
async def generate_cover_letter(
    file: UploadFile = File(...),
    journal_name: Optional[str] = Form(None),
    parsed_rules: Optional[str] = Form(None)
):
    """
    基于模板替换和正则，修改原有 Cover Letter 内容。无需通过大模型。
    """
    import json
    import re
    from docx import Document
    
    rules = {}
    if parsed_rules:
        try:
            parsed_data = json.loads(parsed_rules)
            rules = parsed_data.get("rules", parsed_data)
        except:
            pass

    original_text = ""
    try:
        content = await file.read()
        temp_path = f"temp/cv_{file.filename}"
        with open(temp_path, "wb") as f:
            f.write(content)
        
        if file.filename.lower().endswith('.md') or file.filename.lower().endswith('.txt'):
            with open(temp_path, "r", encoding="utf-8") as f:
                original_text = f.read()
        else:
            doc = Document(temp_path)
            original_text = "\n".join([p.text for p in doc.paragraphs])
    except Exception as e:
        print(f"Failed to read file for cover letter: {e}")
        raise HTTPException(status_code=400, detail=f"无法读取您上传的 Cover Letter 文件，请确保它是 .docx，.md 或 .txt 格式。错误信息: {str(e)}")

    if not original_text.strip():
        raise HTTPException(status_code=400, detail="提取失败：您上传的 Cover Letter 文件似乎是空的。")

    try:
        # 修改逻辑（基于正则匹配替换和末尾添加）
        revised_text = original_text

        # 1. 尝试粗略替换目标期刊名字 (如果在原信中找到了类似 "to [Journal Name]" 的结构)
        if journal_name:
            # Look for common phrases like "submitted to X", "consideration for publication in X"
            # It's hard to accurately guess the OLD journal name via Regex without AI, 
            # so we might look for common placeholders:
            placeholder_pattern = re.compile(r"\[.*?Journal.*?\]|\[.*?Name.*?\]", re.IGNORECASE)
            if placeholder_pattern.search(revised_text):
                 revised_text = placeholder_pattern.sub(journal_name, revised_text)

        # 2. 补齐必填声明 (如果原始文本里没有提到)
        declarations = []
        original_lower = revised_text.lower()
        
        if rules.get("ethics_statement_required") and "ethic" not in original_lower:
            declarations.append("The study was conducted according to the guidelines of the Declaration of Helsinki, and approved by the Institutional Review Board.")
            
        if rules.get("conflict_statement_required") and "conflict" not in original_lower:
            declarations.append("All authors declare that they have no conflicts of interest.")
            
        if rules.get("data_availability_required") and "data" not in original_lower:
            declarations.append("The data underlying this article are available in the manuscript.")

        # 3. 如果需要补充声明，添加到信件末尾（Sincerely / Yours faithfully 之前）
        if declarations:
            declarations_text = "\n\n" + "\n".join(declarations) + "\n\n"
            
            # Find sign-off to insert before it
            signoff_match = re.search(r"\n(Sincerely|Yours faithfully|Best regards|Yours sincerely)", revised_text, re.IGNORECASE)
            
            if signoff_match:
                insert_pos = signoff_match.start()
                revised_text = revised_text[:insert_pos] + declarations_text + revised_text[insert_pos:]
            else:
                revised_text += declarations_text

        # 4. Save the revised text to a new DOCX file
        out_doc = Document()
        for i, paragraph_text in enumerate(revised_text.split('\n')):
            # Simple assumption: empty lines or multiple newlines mean paragraph breaks
            if paragraph_text.strip() == "" and i != len(revised_text.split('\n')) - 1:
                continue # Don't add completely empty paragraphs everywhere
            out_doc.add_paragraph(paragraph_text)
            
        import os
        output_dir = os.path.join(os.getcwd(), "static", "output")
        os.makedirs(output_dir, exist_ok=True)
        cv_filename = f"Revised_{file.filename}"
        cv_path = os.path.join(output_dir, cv_filename)
        out_doc.save(cv_path)

        return {
            "status": "success",
            "message": "Cover Letter 核对与修正完成（本地处理）",
            "data": {
                 "cover_letter": revised_text,
                 "journal_name": journal_name or "Unknown",
                 "download_url": f"/static/output/{cv_filename}"
            },
        }
    except Exception as e:
        err_msg = str(e)
        raise HTTPException(status_code=500, detail=f"修正 Cover Letter 失败: {err_msg}")


@router.post("/format-tables-only")
async def format_tables_only(
    file: UploadFile = File(...),
):
    import os
    import json
    from services.document_splitter import DocumentSplitter
    from services.document_builder import DocumentBuilder

    try:
        os.makedirs("temp", exist_ok=True)
        file_path = f"temp/{file.filename}"
        with open(file_path, "wb") as f:
            f.write(await file.read())

        splitter = DocumentSplitter()
        sections = splitter.parse(file_path)

        out_dir = f"static/output/tables_{os.path.splitext(file.filename)[0]}"
        os.makedirs(out_dir, exist_ok=True)
        builder = DocumentBuilder(out_dir)

        rules = {"font_size_pt": 12, "font_family": "Times New Roman"}

        output_filename = builder.build_tables_only(file.filename, sections, rules)
        
        final_path = f"static/output/{output_filename}"
        import shutil
        shutil.move(os.path.join(out_dir, output_filename), final_path)

        return {
            "status": "success",
            "message": "Table extraction and formatting completed.",
            "data": {
                "filename": output_filename,
                "download_url": f"/static/output/{output_filename}"
            }
        }
    except Exception as e:
        logger.error(f"Error formatting tables: {str(e)}")
        return {"status": "error", "message": str(e)}


@router.post("/format-tables-only")
async def format_tables_only(
    file: UploadFile = File(...),
):
    import os
    import json
    from services.document_splitter import DocumentSplitter
    from services.document_builder import DocumentBuilder

    try:
        os.makedirs("temp", exist_ok=True)
        file_path = f"temp/{file.filename}"
        with open(file_path, "wb") as f:
            f.write(await file.read())

        splitter = DocumentSplitter()
        sections = splitter.parse(file_path)

        out_dir = f"static/output/tables_{os.path.splitext(file.filename)[0]}"
        os.makedirs(out_dir, exist_ok=True)
        builder = DocumentBuilder(out_dir)

        rules = {"font_size_pt": 12, "font_family": "Times New Roman"}

        output_filename = builder.build_tables_only(file.filename, sections, rules)
        
        final_path = f"static/output/{output_filename}"
        import shutil
        shutil.move(os.path.join(out_dir, output_filename), final_path)

        return {
            "status": "success",
            "message": "Table extraction and formatting completed.",
            "data": {
                "filename": output_filename,
                "download_url": f"/static/output/{output_filename}"
            }
        }
    except Exception as e:
        logger.error(f"Error formatting tables: {str(e)}")
        return {"status": "error", "message": str(e)}
